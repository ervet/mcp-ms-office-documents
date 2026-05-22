"""Regression tests for the ``markdown_to_word`` infinite-loop hang.

Background
----------
The bug was observed in EKS production. A py-spy dump showed two
``run_blocking_*`` worker threads pinned at 100 % CPU for 6+ hours each,
both spinning at:

    markdown_to_word (docx_tools/base_docx_tool.py:170)
        process_list_items (docx_tools/helpers.py:264)

Root cause
----------
When a markdown line has **leading whitespace** AND matches a list
pattern on its stripped form (e.g. ``"   1. item"`` strips to
``"1. item"``), the caller in ``markdown_to_word`` dispatches it::

    elif ORDERED_LIST_PATTERN.match(line):           # matches '1. item'
        i, _ = process_list_items(lines, i, doc, True, 0)

…but ``process_list_items`` inspects the **original** line's indent (3
characters → ``current_level = 1``) and breaks immediately on
``current_level != level (0)``, returning ``start_idx`` unchanged. The
caller's outer ``while`` then re-detects the same line, calls again,
gets the same answer — and the loop never advances.

Test mechanism
--------------
``markdown_to_word`` is run in a daemon thread with a wall-clock
``join`` timeout. If the call doesn't return inside the timeout the
test fails with an explicit message. The orphan thread keeps spinning
until the pytest process exits — keep the markdown inputs *small* so
the orphan does not steal too many GIL slots from subsequent tests in
the same pytest session.

Today (bug present) these tests **FAIL** with a clear timeout message.
Once the forward-progress guard is added to ``process_list_items`` they
**PASS** in well under a second.

Debugging tip
-------------
To inspect the live hang interactively, run a single test verbosely::

    pytest -xvs tests/test_docx_list_infinite_loop_regression.py::test_top_level_ordered_list_with_leading_whitespace_does_not_hang

Then, during the 5-second window before the test fails, capture a
py-spy dump in another shell::

    py-spy dump --pid $(pgrep -f pytest)

The dump should mirror the production stack trace (``process_list_items``
called from ``markdown_to_word`` with the line index pinned).
"""

import sys
import threading
from pathlib import Path

import pytest

# Add project root to path for imports
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx import Document  # noqa: E402

from docx_tools.base_docx_tool import markdown_to_word  # noqa: E402
from docx_tools.block_elements import process_list_items  # noqa: E402


# Generous wall-clock budget. Real markdown_to_word on these tiny inputs
# completes in well under 100 ms; 5 s is the "definitely hung" threshold.
_TIMEOUT_SECONDS = 5.0


def _run_in_thread(markdown_text, timeout=_TIMEOUT_SECONDS):
    """Invoke ``markdown_to_word`` in a daemon thread, bounded by wall-clock.

    Args:
        markdown_text: Markdown content fed to ``markdown_to_word``.
        timeout: Max seconds to wait for the call to return.

    Returns:
        Tuple ``(completed, result, error)``:
          * ``completed`` (bool): True if the call returned within ``timeout``.
          * ``result``: whatever ``markdown_to_word`` returned, or ``None``.
          * ``error``: the exception raised by the call, if any.

    The thread is created with ``daemon=True`` so it does not block the
    pytest process from exiting. Python cannot kill a thread that has
    not finished — if the bug is present the thread will keep running
    in the background until the test session ends. Test inputs are kept
    minimal to limit that cost.
    """
    holder = {"result": None, "error": None}

    def target():
        try:
            holder["result"] = markdown_to_word(markdown_text)
        except Exception as exc:  # noqa: BLE001 — re-raised below
            holder["error"] = exc

    t = threading.Thread(
        target=target,
        daemon=True,
        name="markdown-to-word-regression",
    )
    t.start()
    t.join(timeout=timeout)
    return (not t.is_alive(), holder["result"], holder["error"])


def test_top_level_ordered_list_with_leading_whitespace_does_not_hang():
    """REGRESSION: ``'   1. item'`` at the top level previously hung the worker.

    Reproduces the exact stack observed in the EKS py-spy dump.

    Sequence the parser walks (today, buggy):

      1. ``"Intro paragraph before the list."`` — falls into the
         ``else:`` paragraph branch, advances ``i``.
      2. Empty line — handled by the empty-line counter, advances ``i``.
      3. ``"   1. First item..."`` — caller strips it to ``"1. First
         item..."``, matches ``ORDERED_LIST_PATTERN``, calls
         ``process_list_items(..., level=0)``. Inside, the indent of
         the original line gives ``current_level=1``, which does not
         equal ``level=0``, so the function breaks and returns
         ``start_idx`` unchanged. The caller then loops back to step 3
         forever.
    """
    md = (
        "Intro paragraph before the list.\n"
        "\n"
        "   1. First item with 3-space indent\n"
        "   2. Second item with 3-space indent\n"
        "\n"
        "Final paragraph after the list.\n"
    )

    completed, result, error = _run_in_thread(md)

    assert completed, (
        f"markdown_to_word did not complete within {_TIMEOUT_SECONDS:.0f}s "
        f"on input with a top-level ORDERED list that has leading "
        f"whitespace. This is the infinite-loop bug observed in EKS "
        f"production (py-spy dump showed two worker threads pinned in "
        f"markdown_to_word -> process_list_items)."
    )
    if error is not None:
        raise error
    assert result is not None


def test_top_level_unordered_list_with_leading_whitespace_does_not_hang():
    """Same bug as the ordered-list case, with bullets instead of numbers.

    ``UNORDERED_LIST_PATTERN`` matches ``"- bullet"`` on the stripped
    line just like the ordered variant, so the same dispatch path is
    exercised and the same infinite loop occurs in
    ``process_list_items``.
    """
    md = (
        "Some intro paragraph.\n"
        "\n"
        "   - first bullet with 3-space indent\n"
        "   - second bullet with 3-space indent\n"
        "\n"
        "Some outro paragraph.\n"
    )

    completed, result, error = _run_in_thread(md)

    assert completed, (
        f"markdown_to_word did not complete within {_TIMEOUT_SECONDS:.0f}s "
        f"on input with a top-level UNORDERED list that has leading "
        f"whitespace. Same root cause as the ordered-list variant — see "
        f"the module docstring for the full trace."
    )
    if error is not None:
        raise error
    assert result is not None


# ---------------------------------------------------------------------------
# Helper-level unit tests — exercise process_list_items in isolation.
# ---------------------------------------------------------------------------
# These tests do not invoke ``markdown_to_word`` (and therefore do not need
# a threading-based timeout). They call ``process_list_items`` directly and
# assert the *forward-progress* property: regardless of input, the returned
# index must be strictly greater than ``start_idx`` so the caller's outer
# loop is guaranteed to advance.
#
# The bug is that today, when the original line's indent does not match
# the requested ``level`` (which happens when the caller dispatched based
# on the stripped form), ``process_list_items`` breaks on the very first
# iteration and returns ``start_idx`` unchanged. The fix must restore the
# forward-progress invariant.
# ---------------------------------------------------------------------------


def _new_blank_doc():
    """Create a minimal ``docx.Document`` for tests.

    Returns a Document with no template so we don't depend on the local
    ``custom_templates/*.docx`` having particular styles. The bug path
    in ``process_list_items`` breaks BEFORE any ``add_paragraph(style=...)``
    call, so a blank document is sufficient to reproduce the failure.
    The forward-progress fallback (once applied) calls ``add_paragraph()``
    with no style, which is always valid on a blank Document.
    """
    return Document()


def test_process_list_items_advances_on_ordered_line_with_leading_whitespace():
    """REGRESSION (unit-level): process_list_items must return ``i > start_idx``.

    Today this assertion fails because the function returns
    ``start_idx`` unchanged when the original line's indent does not
    match the requested ``level``. The caller (``markdown_to_word``)
    re-dispatches and spins forever — see the threaded tests above.

    Use a debugger here to step through the exact code path:

        breakpoint()  # or set in helpers.py:281 (top of the while loop)
        ...
    """
    lines = ["   1. First item with 3-space indent"]
    start_idx = 0
    doc = _new_blank_doc()

    new_idx, _elements = process_list_items(
        lines, start_idx, doc, is_ordered=True, level=0
    )

    assert new_idx > start_idx, (
        f"process_list_items returned new_idx={new_idx} (== start_idx). "
        f"This violates the forward-progress invariant the caller relies "
        f"on. With this input, markdown_to_word's outer dispatch loop "
        f"would re-detect the same line as a list and spin forever."
    )


def test_process_list_items_advances_on_unordered_line_with_leading_whitespace():
    """Same forward-progress assertion for the bullet-list variant."""
    lines = ["   - bullet with 3-space indent"]
    start_idx = 0
    doc = _new_blank_doc()

    new_idx, _elements = process_list_items(
        lines, start_idx, doc, is_ordered=False, level=0
    )

    assert new_idx > start_idx, (
        f"process_list_items returned new_idx={new_idx} (== start_idx) "
        f"for an unordered list line with leading whitespace. Forward-"
        f"progress invariant violated; see the ordered-list test for the "
        f"full explanation."
    )


def test_process_list_items_advances_on_mid_document_offset():
    """Forward progress must also hold when start_idx is mid-document.

    The bug isn't index-dependent — it's about indent vs level mismatch
    on the very first line the function inspects. We exercise that
    independently of ``start_idx=0`` to make sure no off-by-one creeps
    into the eventual fix.
    """
    lines = [
        "Some earlier content.",
        "",
        "   1. Indented item starts here (start_idx=2)",
    ]
    start_idx = 2
    doc = _new_blank_doc()

    new_idx, _elements = process_list_items(
        lines, start_idx, doc, is_ordered=True, level=0
    )

    assert new_idx > start_idx, (
        f"process_list_items returned new_idx={new_idx} for start_idx={start_idx}; "
        f"the function must always advance past the line it cannot process."
    )


# ---------------------------------------------------------------------------
# Happy-path tests — for comparison with the failing regression tests above.
# ---------------------------------------------------------------------------
# These exercise the code paths the buggy tests *should* end up on once the
# fix is applied: properly-formed lists with no leading whitespace. They
# verify that:
#
#   1. ``process_list_items`` consumes the list lines and emits paragraphs
#      with the correct ``List Number`` / ``List Bullet`` styles (i.e. the
#      ``paragraph = doc.add_paragraph(style=style)`` line at
#      helpers.py:301 is actually executed and works).
#   2. ``markdown_to_word`` runs end-to-end on a document containing both
#      ordered and unordered lists, without hanging.
#
# These tests already pass today (the bug is dormant for well-formed input)
# and must continue to pass after the forward-progress fix lands — they're
# the "control group" for the regression suite.
# ---------------------------------------------------------------------------


def _new_doc_with_default_styles():
    """Create a ``Document`` from the project's default template.

    The default template (``default_templates/default_docx_template.docx``)
    contains every paragraph style the parser uses — ``Heading 1``..``6``,
    ``List Number``, ``List Bullet``, ``Quote``, ``Table Grid``, etc. We
    use it directly so these tests don't depend on whatever template the
    ``custom_templates/`` directory happens to contain locally.
    """
    default = project_root / "default_templates" / "default_docx_template.docx"
    assert default.exists(), f"Default template missing at {default}"
    return Document(str(default))


@pytest.fixture
def use_default_template(monkeypatch):
    """Force ``markdown_to_word`` to load the default template.

    The local ``custom_templates/custom_docx_template.docx`` (BU asset
    bundle) lacks several styles the parser uses on the happy path
    (``Heading 1``, ``List Number``, ``List Bullet``, ``Quote``). Rather
    than mutate the filesystem we patch the ``load_templates`` symbol
    that ``markdown_to_word`` looks up at call time, redirecting it to
    the default template path. The fixture is scoped to a single test
    and ``monkeypatch`` restores the original binding automatically.
    """
    default = project_root / "default_templates" / "default_docx_template.docx"
    assert default.exists(), f"Default template missing at {default}"

    import docx_tools.base_docx_tool as _base
    monkeypatch.setattr(_base, "load_templates", lambda: str(default))
    yield


def test_process_list_items_creates_ordered_list_paragraphs():
    """Happy path at the helper level for ordered (numbered) lists.

    Verifies that:
      * ``new_idx`` advances by the number of list lines consumed,
      * the document gains paragraphs whose style name is ``List Number``,
      * the paragraph text matches the list-item content (the part after
        ``"N. "``, with inline formatting stripped of markers).

    This is exactly what the ``doc.add_paragraph(style=style)`` line at
    [docx_tools/helpers.py:301](docx_tools/helpers.py#L301) is supposed
    to produce on the happy path.
    """
    lines = ["1. First item", "2. Second item", "3. Third item"]
    doc = _new_doc_with_default_styles()
    initial_paragraph_count = len(doc.paragraphs)

    new_idx, _elements = process_list_items(
        lines, 0, doc, is_ordered=True, level=0
    )

    assert new_idx == 3, (
        f"Expected process_list_items to consume all 3 list lines "
        f"(new_idx=3); got new_idx={new_idx}"
    )

    list_paragraphs = [
        p for p in doc.paragraphs[initial_paragraph_count:]
        if p.style.name == "List Number"
    ]
    assert len(list_paragraphs) == 3, (
        f"Expected 3 'List Number'-styled paragraphs; got {len(list_paragraphs)}. "
        f"All new paragraphs: "
        f"{[(p.text, p.style.name) for p in doc.paragraphs[initial_paragraph_count:]]}"
    )
    assert list_paragraphs[0].text == "First item"
    assert list_paragraphs[1].text == "Second item"
    assert list_paragraphs[2].text == "Third item"


def test_process_list_items_creates_unordered_list_paragraphs():
    """Happy path at the helper level for unordered (bullet) lists.

    Same shape as the ordered-list test above, but for the
    ``List Bullet`` style branch.
    """
    lines = ["- alpha", "- beta", "- gamma"]
    doc = _new_doc_with_default_styles()
    initial_paragraph_count = len(doc.paragraphs)

    new_idx, _elements = process_list_items(
        lines, 0, doc, is_ordered=False, level=0
    )

    assert new_idx == 3
    list_paragraphs = [
        p for p in doc.paragraphs[initial_paragraph_count:]
        if p.style.name == "List Bullet"
    ]
    assert len(list_paragraphs) == 3
    assert [p.text for p in list_paragraphs] == ["alpha", "beta", "gamma"]


def test_markdown_to_word_with_valid_lists_completes_promptly(use_default_template):
    """End-to-end happy path: ``markdown_to_word`` with well-formed lists.

    No leading whitespace, so the dispatch path stays on the normal
    branch and ``process_list_items`` consumes every list line. The
    call must return quickly with a non-None result (file path or URL
    depending on the configured upload backend — LOCAL here).

    Use this test alongside the threaded regression tests to confirm
    that the *only* thing different about the buggy input is the
    leading whitespace.
    """
    md = (
        "Introductory paragraph before any lists.\n"
        "\n"
        "1. First ordered item\n"
        "2. Second ordered item\n"
        "3. Third ordered item\n"
        "\n"
        "- First bullet\n"
        "- Second bullet\n"
        "- Third bullet\n"
        "\n"
        "Closing paragraph after the lists.\n"
    )

    completed, result, error = _run_in_thread(md, timeout=10)

    assert completed, (
        "markdown_to_word did not complete on well-formed input — if this "
        "fails, the bug has spread beyond leading-whitespace lines"
    )
    if error is not None:
        raise error
    assert result is not None
    assert isinstance(result, str) and result, (
        f"Expected a non-empty result string; got {result!r}"
    )


def test_process_list_items_handles_nested_unordered_list():
    """Happy path at the helper level for **nested** bullet lists.

    Exercises the recursive call inside ``process_list_items`` (the
    "Look ahead for nested items" inner loop). Nested levels are
    detected via 3-space indent increments — ``level=0`` for the top
    list, ``level=1`` for sub-items.

    The input:

        - Main item 1
           - Sub item 1.1     ← 3-space indent → level 1
           - Sub item 1.2     ← still level 1
        - Main item 2          ← back to level 0
           - Sub item 2.1     ← level 1 again

    Expected document state after the call:

        ┌────────────────┬──────────────────┐
        │ Paragraph text │ Style            │
        ├────────────────┼──────────────────┤
        │ Main item 1    │ List Bullet      │
        │ Sub item 1.1   │ List Bullet 2    │
        │ Sub item 1.2   │ List Bullet 2    │
        │ Main item 2    │ List Bullet      │
        │ Sub item 2.1   │ List Bullet 2    │
        └────────────────┴──────────────────┘

    Useful for stepping a debugger through the recursive descent:
    drop a ``breakpoint()`` just before the recursive call site in
    ``process_list_items`` (the ``i, nested = process_list_items(...)``
    line in the inner lookahead loop) and watch ``level`` increase
    from 0 → 1 then return.
    """
    lines = [
        "- Main item 1",
        "   - Sub item 1.1",
        "   - Sub item 1.2",
        "- Main item 2",
        "   - Sub item 2.1",
    ]
    doc = _new_doc_with_default_styles()
    initial_paragraph_count = len(doc.paragraphs)

    markdown_cnt = """   - Main item 1
     - Sub item 1.1
     - Sub item 1.2
- Main item 2
   - Sub item 2.1
"""
    # markdown_to_word(markdown_cnt, file_name="MZA-markdown-regression.docx")

    new_idx, _elements = process_list_items(
        lines, 0, doc, is_ordered=False, level=0
    )

    # All 5 lines consumed (no trailing empty line in this slice).
    assert new_idx == len(lines), (
        f"Expected to consume all {len(lines)} lines; got new_idx={new_idx}"
    )

    new_paragraphs = doc.paragraphs[initial_paragraph_count:]
    text_and_style = [(p.text, p.style.name) for p in new_paragraphs]

    expected = [
        ("Main item 1",  "List Bullet"),
        ("Sub item 1.1", "List Bullet 2"),
        ("Sub item 1.2", "List Bullet 2"),
        ("Main item 2",  "List Bullet"),
        ("Sub item 2.1", "List Bullet 2"),
    ]
    assert text_and_style == expected, (
        f"Nested list paragraphs do not match expected output.\n"
        f"  expected: {expected}\n"
        f"  got:      {text_and_style}"
    )


def test_markdown_to_word_with_nested_lists_completes_promptly(use_default_template):
    """End-to-end happy path for nested lists, mirroring the markdown in
    the user-facing tools.

    Same content as ``test_process_list_items_handles_nested_unordered_list``
    but driven through the public ``markdown_to_word`` entry point. This
    exercises the full pipeline: outer ``markdown_to_word`` loop →
    ``UNORDERED_LIST_PATTERN`` match → ``process_list_items(... level=0)``
    → recursive call at level 1 → return → continue parsing.
    """
    md = (
        "- Main item 1\n"
        "   - Sub item 1.1\n"
        "   - Sub item 1.2\n"
        "- Main item 2\n"
        "   - Sub item 2.1\n"
    )

    completed, result, error = _run_in_thread(md, timeout=10)

    assert completed, (
        "markdown_to_word did not complete on a well-formed nested bullet "
        "list — recursive descent in process_list_items may be broken"
    )
    if error is not None:
        raise error
    assert result is not None


# ---------------------------------------------------------------------------
# Forward-progress verification for additional edge cases.
# ---------------------------------------------------------------------------
# These tests cover inputs that exercise the various entry conditions of
# ``process_list_items`` to make sure the forward-progress invariant
# (``returned_idx > start_idx``) holds for every shape of pathological
# input we can think of. They also confirm the recursive descent and the
# ``return_elements`` mode still work after the fix is applied.
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "indent_spaces",
    [3, 4, 6, 9, 12],
    ids=lambda n: f"indent={n}",
)
def test_process_list_items_advances_on_various_indent_widths(indent_spaces):
    """Forward progress must hold for any leading-whitespace width, not
    just the canonical 3-space indent."""
    line = (" " * indent_spaces) + "1. Indented item"
    doc = _new_blank_doc()
    new_idx, _ = process_list_items([line], 0, doc, is_ordered=True, level=0)
    assert new_idx > 0, (
        f"process_list_items did not advance for indent={indent_spaces} spaces"
    )


def test_process_list_items_advances_on_tab_indent():
    """Tabs in the leading whitespace must not cause a hang either.

    A tab character contributes 1 to the indent count (Python's
    ``str.lstrip`` strips all whitespace including tabs), so a single
    leading tab yields ``current_level = 0`` and the line would
    actually pass the indent check. Multiple tabs (or tab+space)
    produce a non-zero indent that mismatches level=0 and previously
    triggered the hang. Either way, the function must advance.
    """
    lines = ["\t\t1. Tab-indented item"]
    doc = _new_blank_doc()
    new_idx, _ = process_list_items(lines, 0, doc, is_ordered=True, level=0)
    assert new_idx > 0


def test_process_list_items_advances_on_mixed_valid_and_invalid_input():
    """A valid list followed by an indented list line is handled by the
    existing nested-list recursion path, not the forward-progress guard.

    Trace for these three lines:
      i=0  "1. Valid item one"        consumed at level 0
      i=1  "2. Valid item two"        consumed at level 0; inner lookahead
                                       sees line i=2 with indent=3 (next_level=1
                                       > level=0) and recurses
      i=2  "   3. Indented..."        consumed at level 1 by the recursive
                                       call (current_level=1 == level=1)

    All three lines are consumed (``new_idx == 3``). The
    forward-progress guard is a no-op here. The infinite-loop bug only
    triggered when the indented line appeared at the TOP level (no
    valid list item directly above it to anchor recursion). Those
    cases are covered by the earlier tests.
    """
    lines = [
        "1. Valid item one",
        "2. Valid item two",
        "   3. Indented after valid list (treated as nested)",
    ]
    doc = _new_doc_with_default_styles()
    new_idx, _ = process_list_items(lines, 0, doc, is_ordered=True, level=0)

    assert new_idx == 3, (
        f"All 3 lines should be consumed (last one as a nested item); "
        f"got new_idx={new_idx}"
    )


def test_markdown_to_word_with_mixed_valid_and_invalid_lists(use_default_template):
    """End-to-end: the user's invalid input does NOT prevent valid content
    from being rendered. The bug-trigger lines become plain paragraphs;
    the rest of the document is processed normally."""
    md = (
        "Intro.\n"
        "\n"
        "   1. Indented (bug trigger)\n"
        "   2. Another indented\n"
        "\n"
        "1. First valid ordered\n"
        "2. Second valid ordered\n"
        "\n"
        "   - Indented bullet (bug trigger)\n"
        "\n"
        "- First valid bullet\n"
        "- Second valid bullet\n"
        "\n"
        "Outro.\n"
    )

    completed, result, error = _run_in_thread(md, timeout=10)

    assert completed, "markdown_to_word did not complete on mixed valid+invalid"
    if error is not None:
        raise error
    assert result is not None


def test_process_list_items_handles_empty_lines_input():
    """Pure empty/whitespace input must terminate immediately.

    The function is documented to consume list items; if there are no
    list items at start_idx, it must still return ``new_idx > start_idx``
    so the caller is not stuck.
    """
    lines = [""]
    doc = _new_blank_doc()
    new_idx, _ = process_list_items(lines, 0, doc, is_ordered=True, level=0)
    assert new_idx > 0


def test_process_list_items_handles_non_list_input():
    """A plain paragraph line at start_idx (not a list pattern at all)
    must still result in forward progress.

    This case can be reached if a future caller dispatches incorrectly,
    or if the input contains a list pattern only on the stripped form
    but not in the cleaned line (e.g. unusual whitespace). The guard
    catches every "no-progress" scenario uniformly.
    """
    lines = ["This is not a list line."]
    doc = _new_blank_doc()
    new_idx, _ = process_list_items(lines, 0, doc, is_ordered=True, level=0)
    assert new_idx > 0


def test_process_list_items_recursive_call_unaffected_by_guard():
    """The forward-progress guard must be a no-op for recursive calls.

    When ``process_list_items`` recurses for nested items it passes the
    nested level explicitly. The first iteration's indent matches that
    level by construction, so the function consumes at least one item
    and the guard at the end never fires. This test asserts that
    behaviour is unchanged: a flat list with a nested sub-list produces
    the expected paragraphs and styles, exactly as before the fix.
    """
    lines = [
        "- Outer item",
        "   - Nested item",
    ]
    doc = _new_doc_with_default_styles()
    initial = len(doc.paragraphs)

    new_idx, _ = process_list_items(lines, 0, doc, is_ordered=False, level=0)

    assert new_idx == 2
    new_paragraphs = doc.paragraphs[initial:]
    text_and_style = [(p.text, p.style.name) for p in new_paragraphs]
    assert text_and_style == [
        ("Outer item", "List Bullet"),
        ("Nested item", "List Bullet 2"),
    ], (
        "Recursive descent for nested items behaves differently after the "
        "forward-progress guard was added — the guard should be a no-op "
        f"in the recursion path. Got: {text_and_style}"
    )


def test_process_list_items_return_elements_mode_works_with_guard():
    """The ``return_elements=True`` mode (used by template-placeholder
    rendering in ``dynamic_docx_tools``) must keep working — both on the
    happy path and when the guard fires.

    On the happy path the guard is a no-op. When the guard fires for a
    bug-trigger line, the fallback paragraph it creates is also detached
    from the body and appended to ``elements``, matching the rest of the
    function's contract.
    """
    # Happy path with return_elements=True
    happy_lines = ["1. Captured item"]
    doc = _new_doc_with_default_styles()
    new_idx, elements = process_list_items(
        happy_lines, 0, doc, is_ordered=True, level=0, return_elements=True
    )
    assert new_idx == 1
    assert elements is not None and len(elements) == 1, (
        f"Expected one element returned on happy path; got {elements}"
    )

    # Guard path with return_elements=True — bug-trigger line should
    # produce one fallback element and still advance.
    guarded_lines = ["   1. Indented (bug trigger)"]
    doc2 = _new_doc_with_default_styles()
    new_idx2, elements2 = process_list_items(
        guarded_lines, 0, doc2, is_ordered=True, level=0, return_elements=True
    )
    assert new_idx2 == 1
    assert elements2 is not None and len(elements2) == 1, (
        "Forward-progress guard with return_elements=True did not append "
        f"the fallback paragraph to elements; got {elements2}"
    )


def test_markdown_to_word_with_only_indented_list_completes(use_default_template):
    """Worst-case input: the entire document is bug-trigger lines.

    Even when every list line in the document has leading whitespace,
    ``markdown_to_word`` must terminate. Each line gets rendered as a
    plain paragraph via the guard."""
    md = (
        "   1. Indented one\n"
        "   2. Indented two\n"
        "   3. Indented three\n"
    )

    completed, result, error = _run_in_thread(md, timeout=10)
    assert completed
    if error is not None:
        raise error
    assert result is not None


def test_markdown_to_word_on_empty_input(use_default_template):
    """Empty markdown must not hang and must produce a (possibly empty)
    document."""
    completed, result, error = _run_in_thread("", timeout=10)
    assert completed
    if error is not None:
        raise error
    assert result is not None


def test_markdown_to_word_on_whitespace_only_input(use_default_template):
    """Whitespace-only input (newlines and spaces) must terminate too."""
    completed, result, error = _run_in_thread("\n   \n\n   \n", timeout=10)
    assert completed
    if error is not None:
        raise error
    assert result is not None
