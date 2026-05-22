"""Document-level features: header/footer, TOC, template loading."""
import logging
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from template_utils import find_docx_template
from .patterns import _PAGE_TOKEN_RE
logger = logging.getLogger(__name__)
def load_templates():
    """Resolve Word template path from custom/default template directories.
    Returns absolute path as string or None if not found.
    """
    path = find_docx_template()
    if path:
        logger.debug(f"Using Word template: {path}")
    else:
        logger.warning("No Word template found, will create a blank document")
    return path
# ---------------------------------------------------------------------------
# Header / footer / Word fields
# ---------------------------------------------------------------------------
def _add_field(paragraph, field_code):
    """Insert a Word field (PAGE, NUMPAGES, etc.) into a paragraph."""
    for fld_type, text in [('begin', None), (None, field_code), ('end', None)]:
        run = paragraph.add_run()
        if fld_type:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), fld_type)
            run._r.append(fld)
        else:
            elem = OxmlElement('w:instrText')
            elem.set(qn('xml:space'), 'preserve')
            elem.text = f' {text} '
            run._r.append(elem)
def set_header_footer(doc, text, kind='header'):
    """Set document header or footer text.
    Iterates over **all** document sections.  For each section the default
    header/footer is updated, and - when the section uses a different first-page
    header/footer - that variant is updated as well.
    Pre-existing paragraph formatting (alignment, style) from the template is
    preserved; only run content is replaced.
    Args:
        doc: The Word document.
        text: Content string.  Use ``{page}`` / ``{pages}`` for field tokens.
        kind: ``'header'`` or ``'footer'``.
    """
    _TOKEN_MAP = {'{page}': 'PAGE', '{pages}': 'NUMPAGES'}
    def _fill_paragraph(p, content):
        """Clear existing runs/fields and write *content* into paragraph *p*."""
        existing_alignment = p.alignment
        for child in list(p._p):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('r', 'hyperlink', 'fldSimple'):
                p._p.remove(child)
        for part in _PAGE_TOKEN_RE.split(content):
            if part in _TOKEN_MAP:
                _add_field(p, _TOKEN_MAP[part])
            elif part:
                p.add_run(part)
        p.alignment = existing_alignment if existing_alignment is not None else WD_ALIGN_PARAGRAPH.CENTER
    def _update_part(section_part):
        """Update a single header or footer part."""
        section_part.is_linked_to_previous = False
        if section_part.paragraphs:
            _fill_paragraph(section_part.paragraphs[0], text)
        else:
            p = section_part.add_paragraph()
            _fill_paragraph(p, text)
    for section in doc.sections:
        _update_part(getattr(section, kind))
        if section.different_first_page_header_footer:
            first_kind = f'first_page_{kind}'
            first_part = getattr(section, first_kind, None)
            if first_part is not None:
                _update_part(first_part)
        even_kind = f'even_page_{kind}'
        even_part = getattr(section, even_kind, None)
        if even_part is not None and doc.settings.element.find(qn('w:evenAndOddHeaders')) is not None:
            _update_part(even_part)


# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------


def add_toc(doc):
    """Insert a Table of Contents field.
    The TOC is based on Heading styles 1-3 and will update when the document
    is opened in Word.
    """
    doc.add_heading('Table of Contents', level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld)
    run = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)
    run = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld)
    p.add_run('[Table of Contents - open in Word and press F9 to update]')
    run = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'end')
    run._r.append(fld)
    doc.add_page_break()
    uf = OxmlElement('w:updateFields')
    uf.set(qn('w:val'), 'true')
    doc.settings.element.append(uf)
