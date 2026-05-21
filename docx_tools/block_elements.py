"""Block-level markdown elements: tables, lists, images, alignment, horizontal lines."""
import logging
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .patterns import (
    ORDERED_LIST_PATTERN,
    UNORDERED_LIST_PATTERN,
    ORDERED_LIST_CAPTURE_PATTERN,
    UNORDERED_LIST_CAPTURE_PATTERN,
    _ALIGN_INLINE_RE,
    _ALIGN_OPEN_RE,
    _ALIGN_CLOSE_RE,
)
from .inline_formatting import parse_inline_formatting
logger = logging.getLogger(__name__)
ALIGNMENT_MAP = {
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'left': WD_ALIGN_PARAGRAPH.LEFT,
}
# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------
def parse_table(lines, start_idx):
    """Parse markdown table and return the table data and next line index."""
    table_lines = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            i += 1
        else:
            break
    if len(table_lines) < 2:
        return None, start_idx + 1
    table_data = []
    for line in table_lines:
        if '---' in line or ':-:' in line or ':--' in line or '--:' in line:
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        table_data.append(cells)
    return table_data, i
def add_table_to_doc(table_data, doc):
    """Add table data to Word document.
    Returns the created ``Table`` object, or ``None`` when the table could
    not be created (empty data or exception).
    """
    if not table_data:
        return None
    rows = len(table_data)
    cols = max(len(row) for row in table_data) if table_data else 0
    try:
        word_table = doc.add_table(rows=rows, cols=cols)
        word_table.style = 'Table Grid'
    except Exception as e:
        logger.warning("Failed to create table with 'Table Grid' style, using default: %s", e)
        try:
            word_table = doc.add_table(rows=rows, cols=cols)
        except Exception as e2:
            logger.error("Failed to create table: %s", e2, exc_info=True)
            return None
    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                try:
                    cell = word_table.cell(i, j)
                    if cell.paragraphs:
                        cell.paragraphs[0].clear()
                    parse_inline_formatting(cell_text, cell.paragraphs[0])
                except Exception as e:
                    logger.warning("Failed to populate table cell [%d, %d]: %s", i, j, e)
    return word_table
# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------
def process_list_items(lines, start_idx, doc, is_ordered=False, level=0,
                       return_elements=False):
    """Process markdown list items with proper Word numbering.
    When *return_elements* is True the created paragraph XML elements are
    removed from the document body and returned so the caller can re-insert
    them elsewhere (used by the template placeholder machinery).
    Returns:
        Tuple of (next_line_index, list_of_elements | None).
    """
    bullet_styles = ['List Bullet', 'List Bullet 2', 'List Bullet 3']
    number_styles = ['List Number', 'List Number 2', 'List Number 3']
    style_array = number_styles if is_ordered else bullet_styles
    style = style_array[min(level, len(style_array) - 1)]
    elements = [] if return_elements else None
    i = start_idx
    n = len(lines)
    list_capture_pattern = (
        ORDERED_LIST_CAPTURE_PATTERN if is_ordered else UNORDERED_LIST_CAPTURE_PATTERN
    )
    while i < n:
        original_line = lines[i]
        stripped_left = original_line.lstrip()
        indent = len(original_line) - len(stripped_left)
        line = stripped_left.rstrip()
        current_level = indent // 3
        if current_level != level:
            break
        list_match = list_capture_pattern.match(line)
        if not list_match:
            break
        paragraph = doc.add_paragraph(style=style)
        parse_inline_formatting(list_match.group(1), paragraph)
        if return_elements:
            elements.append(paragraph._p)
            doc._body._body.remove(paragraph._p)
        i += 1
        # Look ahead for nested items
        while i < n:
            next_original = lines[i]
            next_stripped_left = next_original.lstrip()
            next_line = next_stripped_left.rstrip()
            if not next_line:
                i += 1
                continue
            next_indent = len(next_original) - len(next_stripped_left)
            next_level = next_indent // 3
            if next_level > level:
                is_nested_ordered = bool(ORDERED_LIST_PATTERN.match(next_line))
                is_nested_unordered = bool(UNORDERED_LIST_PATTERN.match(next_line))
                if is_nested_ordered or is_nested_unordered:
                    i, nested = process_list_items(
                        lines, i, doc, is_nested_ordered, next_level, return_elements
                    )
                    if return_elements and nested:
                        elements.extend(nested)
                else:
                    break
            else:
                break
    # Forward-progress guarantee
    if i == start_idx:
        original_line = lines[start_idx]
        stripped_line = original_line.strip()
        logger.warning(
            "process_list_items: no progress at line %d (%r) for level=%d; "
            "rendering as a plain paragraph to guarantee forward progress",
            start_idx, stripped_line, level,
        )
        paragraph = doc.add_paragraph()
        parse_inline_formatting(stripped_line, paragraph)
        if return_elements:
            elements.append(paragraph._p)
            doc._body._body.remove(paragraph._p)
        i = start_idx + 1
    return i, elements
# ---------------------------------------------------------------------------
# Page break / horizontal line
# ---------------------------------------------------------------------------
def add_horizontal_line(doc):
    """Add a visual horizontal line (thin border) to the document."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p
# ---------------------------------------------------------------------------
# Images
# ---------------------------------------------------------------------------
def add_image_to_doc(doc, url, alt_text, max_width_inches=None):
    """Add an image from a URL to the document.
    Downloads the image and inserts it.  On failure inserts an error
    placeholder paragraph instead.
    """
    try:
        from pptx_tools.image_utils import download_image
        if max_width_inches is None:
            try:
                sec = doc.sections[-1]
                max_width_inches = (sec.page_width - sec.left_margin - sec.right_margin) / 914400
            except Exception:
                max_width_inches = 5.5
        image_stream, _ = download_image(url)
        doc.add_picture(image_stream, width=Inches(max_width_inches))
        if alt_text:
            caption = doc.add_paragraph()
            caption.add_run(alt_text).italic = True
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.warning("Failed to add image from '%s': %s", url, e)
        doc.add_paragraph().add_run(f"[Image could not be loaded: {url}]")
# ---------------------------------------------------------------------------
# Text alignment
# ---------------------------------------------------------------------------
def detect_alignment(line):
    """Detect an alignment tag (inline *or* block-open) on *line*.
    Returns ``(inner_text, alignment)`` for an inline tag,
    ``(None, alignment)`` for a block-open tag, or ``None`` if no match.
    """
    m = _ALIGN_INLINE_RE.match(line)
    if m:
        if m.group(1) is not None:
            return m.group(1).strip(), WD_ALIGN_PARAGRAPH.CENTER
        return m.group(3).strip(), ALIGNMENT_MAP.get(m.group(2).lower(), WD_ALIGN_PARAGRAPH.LEFT)
    m = _ALIGN_OPEN_RE.match(line)
    if m:
        align = ALIGNMENT_MAP.get((m.group(1) or 'center').lower(), WD_ALIGN_PARAGRAPH.CENTER)
        return None, align
    return None
def process_alignment_block(lines, start_idx, doc, alignment, return_elements=False):
    """Process lines inside a multi-line alignment block."""
    elements = [] if return_elements else None
    i = start_idx
    while i < len(lines):
        stripped = lines[i].strip()
        if _ALIGN_CLOSE_RE.match(stripped):
            i += 1
            break
        if not stripped:
            i += 1
            continue
        para = doc.add_paragraph()
        para.alignment = alignment
        parse_inline_formatting(stripped, para)
        if return_elements:
            elements.append(para._p)
            doc._body._body.remove(para._p)
        i += 1
    return i, elements
